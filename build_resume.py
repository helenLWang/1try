#!/usr/bin/env python3
"""Build Le Wang resume in Image-1 TPM Word format. Body copy kept verbatim."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x00, 0x33, 0x99)
BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Times New Roman"
BODY = 10.5
HEADER = 13
NAME = 22
COMPANY = 11


def _set_run_fonts(run, font=FONT):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rFonts.set(qn("w:eastAsia"), font)
    rFonts.set(qn("w:cs"), font)


def style_run(run, size=BODY, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.color.rgb = color
    _set_run_fonts(run)


def set_spacing(p, before=0, after=0, line=240):
    pPr = p._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(int(before * 20)))
    spacing.set(qn("w:after"), str(int(after * 20)))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:beforeAutospacing"), "0")
    spacing.set(qn("w:afterAutospacing"), "0")


def add_bottom_border(paragraph, sz="12", color="000000"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.42)
    section.bottom_margin = Inches(0.38)

    styles = doc.styles["Normal"]
    styles.font.name = FONT
    styles.font.size = Pt(BODY)
    styles.font.color.rgb = BLACK
    rPr = styles.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rFonts.set(qn("w:eastAsia"), FONT)
    rFonts.set(qn("w:cs"), FONT)

    pPr = styles.element.get_or_add_pPr()
    contextual = OxmlElement("w:contextualSpacing")
    pPr.append(contextual)


def tab_pos(doc):
    s = doc.sections[0]
    return s.page_width - s.left_margin - s.right_margin


def new_p(doc, before=0, after=0, line=240, align="left"):
    p = doc.add_paragraph()
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    set_spacing(p, before=before, after=after, line=line)
    p.paragraph_format.tab_stops.add_tab_stop(tab_pos(doc), WD_TAB_ALIGNMENT.RIGHT)
    p.paragraph_format.left_indent = Inches(0)
    p.paragraph_format.right_indent = Inches(0)
    p.paragraph_format.first_line_indent = Inches(0)
    return p


def add_text(p, text, size=BODY, bold=False, italic=False, color=BLACK):
    run = p.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return run


def section_header(doc, title, first=False):
    p = new_p(doc, before=8 if first else 11, after=2, line=240)
    add_text(p, title, size=HEADER, bold=True)
    add_bottom_border(p, sz="12")
    return p


def two_col(
    doc,
    left,
    right,
    left_size=COMPANY,
    left_bold=True,
    left_italic=False,
    left_color=NAVY,
    right_size=BODY,
    right_bold=False,
    right_italic=False,
    before=0,
    after=0,
    line=230,
):
    p = new_p(doc, before=before, after=after, line=line)
    add_text(p, left, size=left_size, bold=left_bold, italic=left_italic, color=left_color)
    if right:
        add_text(p, "\t")
        add_text(p, right, size=right_size, bold=right_bold, italic=right_italic, color=BLACK)
    return p


def segs(text, bold_spans):
    out = []
    cursor = 0
    for span in bold_spans:
        idx = text.find(span, cursor)
        if idx < 0:
            raise ValueError(f"bold span not found: {span!r}")
        if idx > cursor:
            out.append((text[cursor:idx], False, False))
        out.append((span, True, False))
        cursor = idx + len(span)
    if cursor < len(text):
        out.append((text[cursor:], False, False))
    if not out:
        out.append((text, False, False))
    return out


def bullet(doc, title, body_segments, size=BODY):
    p = new_p(doc, before=0.5, after=0.3, line=222, align="left")
    p.paragraph_format.left_indent = Inches(0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    add_text(p, "•  ", size=size)
    add_text(p, f"{title}: ", size=size, bold=True)
    for text, bold, italic in body_segments:
        add_text(p, text, size=size, bold=bold, italic=italic)
    return p


def skill_line(doc, label, rest):
    p = new_p(doc, before=0.4, after=0.4, line=222)
    add_text(p, f"{label}: ", size=BODY, bold=True)
    add_text(p, rest, size=BODY, bold=False)
    return p


def italic_blurb(doc, text):
    p = new_p(doc, before=0, after=0.8, line=222)
    add_text(p, text, size=BODY, italic=True)
    return p


def build():
    doc = Document()
    configure_doc(doc)

    p = new_p(doc, before=0, after=1, line=240, align="center")
    add_text(p, "Le Wang", size=24, bold=True)

    p = new_p(doc, before=0, after=1, line=220, align="center")
    add_text(p, "Pittsburgh, PA  |  lew2@andrew.cmu.edu  |  +1 (412)-812-3916", size=11)

    # ----- Education -----
    section_header(doc, "Education", first=True)

    two_col(
        doc,
        "Carnegie Mellon University, Heinz College",
        "Aug 2026 – Dec 2027",
        before=3,
        after=0,
    )
    p = new_p(doc, before=0, after=0, line=230)
    add_text(
        p,
        "Master of Information Systems Management, Business Intelligence and Data Analytics (",
        size=BODY,
        italic=True,
    )
    add_text(p, "Merit Scholarship", size=BODY, bold=True, italic=True)
    add_text(p, ")", size=BODY, italic=True)
    p = new_p(doc, before=0.4, after=1.5, line=222)
    add_text(p, "Relevant Coursework: ", size=BODY, bold=True)
    add_text(
        p,
        "Agentic AI Development, Machine Learning in Production, A/B Testing, Product Management",
        size=BODY,
    )

    two_col(
        doc,
        "Guangdong University of Finance and Economics",
        "Sep 2020 – Jun 2024",
        before=2.5,
        after=0,
    )
    two_col(
        doc,
        "Bachelor of Business Administration, Certified Accountant (ACCA)",
        "",
        left_size=BODY,
        left_bold=False,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )

    # ----- Experience -----
    section_header(doc, "Experience")

    # 1) DeepWisdom — AIPM
    two_col(doc, "DeepWisdom", "Mar 2026 – Jul 2026", before=7, after=0)
    two_col(
        doc,
        "AI Product Manager Intern",
        "Remote | San Francisco, CA",
        left_size=BODY,
        left_bold=True,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )
    italic_blurb(doc, "$700M ARR AI app-building startup serving North America")
    bullet(
        doc,
        "Feature Launch",
        segs(
            "Drove a 5.2% paid conversion lift and a 9% new user conversion rate lift by owning PRD-to-launch for Excel/PPT generation and Google Workspace integrations features.",
            ["Drove a 5.2% paid conversion lift", "9% new user conversion rate lift"],
        ),
    )
    bullet(
        doc,
        "Engagement Growth",
        segs(
            "Increased average session length by 12% in 3 months by launching a personalized share feature using Claude, iterating through A/B testing, competitive analysis, and user feedback.",
            ["Increased average session length by 12%"],
        ),
    )
    bullet(
        doc,
        "Agentic Workflow",
        segs(
            "Reduced research-to-prototype cycle by 50% by building a PRD review agent that auto-screens user needs analysis, feasibility analysis and prioritization analysis before building.",
            ["Reduced research-to-prototype cycle by 50%"],
        ),
    )
    bullet(
        doc,
        "Paid Acquisition",
        segs(
            "Improved Google Ads ROAS to 2.0x by segmenting users into 4 tiers by purchase speed and ARPU, monitoring Day 1-8 launch performance of a $7K/day Google Ads campaign across high-value markets.",
            ["Improved Google Ads ROAS to 2.0x"],
        ),
    )
    bullet(
        doc,
        "Referral Growth",
        segs(
            "Increased referral-driven signups by 8% by launching Share module, enabling frictionless project collaboration and turned users into advocates.",
            ["Increased referral-driven signups by 8%"],
        ),
    )

    # 2) Insight Solutions — TPM
    two_col(doc, "Insight Solutions", "Nov 2026 – Mar 2026", before=7, after=0)
    two_col(
        doc,
        "Technical Product Manager",
        "Shenzhen, China",
        left_size=BODY,
        left_bold=True,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )
    italic_blurb(doc, "B2B supplier-sourcing platform for U.S. buyers")
    bullet(
        doc,
        "Product Discovery",
        segs(
            "Reduced supplier screening from weeks/months of manual outreach to a <1-minute profile review flow by defining PRD scope for supplier search, 20+ filters, anonymized previews, and 2-4 supplier comparison.",
            ["<1-minute", "20+ filters", "2-4 supplier comparison"],
        ),
    )
    bullet(
        doc,
        "Monetization Design",
        segs(
            "Designed tiered monetization for Starter, Standard, and Pro users by mapping access rules.",
            [],
        ),
    )
    bullet(
        doc,
        "Launch Roadmap",
        segs(
            "Set clear launch targets and success metrics for the supplier platform by translating the roadmap into 60+ commodity/process categories, 80+ supplier data fields, 1,200 Phase 1 profiles, and 3,000 Phase 2 profiles.",
            ["60+ commodity/process categories", "80+ supplier data fields", "1,200 Phase 1 profiles", "3,000 Phase 2 profiles"],
        ),
    )

    # 3) Qrent — moved into Experience, TPM
    two_col(doc, "Qrent", "Nov 2025 – Mar 2026", before=7, after=0)
    two_col(
        doc,
        "Technical Product Manager",
        "",
        left_size=BODY,
        left_bold=True,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )
    italic_blurb(doc, "AI rental platform for international students in Sydney")
    bullet(
        doc,
        "Search Workflow",
        segs(
            "Cut apartment search time from 8 weeks to 6 by shipping a 4-step AI workflow (clarify – search – rank - recommend) with a RAG assistant that uses function calling to explain scoring logic to users.",
            ["8 weeks to 6"],
        ),
    )
    bullet(
        doc,
        "Ranking Agent",
        segs(
            "Built a few-shot prompt scoring agent and web scraper to rank listings by budget and commute, feeding results into the RAG assistant via OpenAI and Claude APIs.",
            ["Built a few-shot prompt scoring agent"],
        ),
    )
    bullet(
        doc,
        "Data-driven MVP",
        segs(
            "Hit 1,100 active users and 8,000 clicks in 6 months by shipping a full-stack MVP (Flask + MySQL + JS) with AI scoring, rental Q&A, and smart filtering.",
            ["1,100 active users", "8,000 clicks"],
        ),
    )

    # 4) EY — original non-PM title kept
    two_col(doc, "EY", "Oct 2024 – Oct 2025", before=7, after=0)
    two_col(
        doc,
        "Associate 2, Audit & Assurance",
        "Shenzhen, China",
        left_size=BODY,
        left_bold=True,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )
    bullet(
        doc,
        "Audit Automation",
        segs(
            "Identified a manual-review bottleneck across a 5-person audit team, then built a Power BI dashboard (Power Query + DAX measures + Claude Code) automating 95% of spot-checks.",
            ["95%"],
        ),
    )
    bullet(
        doc,
        "Root-Cause Analysis",
        segs(
            "Led root-cause analysis across 10K+ transactions (grouped and drilled down by error type) and won stakeholder approval for a $50K fix to a recurring error.",
            ["10K+ transactions", "$50K"],
        ),
    )
    bullet(
        doc,
        "Risk Intelligence",
        segs(
            "Cut manual review time by 70% by building an AI risk agent for a client's CFO office, parsing 10M+ monthly transactions and auto-flagging top-5 anomalies, ensuring auditable outcomes.",
            ["70%", "10M+ monthly transactions"],
        ),
    )

    # 5) NielsenIQ — TPM Intern
    two_col(doc, "NielsenIQ", "Nov 2023 – Mar 2024", before=7, after=0)
    two_col(
        doc,
        "Technical Product Manager Intern",
        "Guangzhou, China",
        left_size=BODY,
        left_bold=True,
        left_italic=True,
        left_color=BLACK,
        before=0,
        after=0,
    )
    italic_blurb(doc, "Top1 Market Analysis Company in North America")
    bullet(
        doc,
        "API Integration",
        segs(
            "Scoped a REST API connector that cut cross-system data connection from 2 weeks to 5 day, and reduced AI forecast variance 10% for P&G, influencing $780K in commercialization decisions.",
            ["2 weeks to 5 day", "10%", "$780K"],
        ),
    )

    # ----- Skills -----
    section_header(doc, "Skills")
    skill_line(
        doc,
        "Product & Delivery",
        "A/B Testing, Agile Development, Roadmap Prioritization, Stakeholder Alignment, Iterative Deployment, Data Analysis, Figma (UX/UI), Cursor, PowerPoint, Excel, PowerBI",
    )
    skill_line(
        doc,
        "AI & Engineering",
        "AI Agents & Skills building, AI Workflow Design, AI Demo Prototyping, API, Prompt Engineering, Python, SQL, Flask, MySQL, LLM APIs (OpenAI, Claude), JIRA, GIT",
    )

    out = "/workspace/Le_Wang_AI_PM_Resume.docx"
    doc.save(out)
    print("Wrote", out)


if __name__ == "__main__":
    build()
